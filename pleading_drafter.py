"""Utilities for generating pleadings from templates."""

from __future__ import annotations
import pathlib as _pl
import re
import io
from typing import Dict, Tuple, Any

from config import logger, GEMINI_MODEL_DEFAULT, get_gemini_model
from ai_utils import DEFAULT_GEMINI_SAFETY_SETTINGS, _gemini_generate_content_with_retry_and_tokens
try:
    from docx import Document
    from google.generativeai import types as genai_types  # type: ignore
except Exception:
    Document = None
    genai_types = None  # type: ignore


def load_template(path: _pl.Path) -> str:
    """Return template text from a Markdown or DOCX file."""
    if not path.exists():
        raise FileNotFoundError(path)

    if path.suffix.lower() == ".md":
        return path.read_text(encoding="utf-8")
    if path.suffix.lower() == ".docx" and Document:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    raise ValueError(f"Unsupported template type: {path.suffix}")


def _render_markdown_to_docx(text: str, output: io.BytesIO) -> None:
    if not Document:
        raise ImportError("python-docx library required for DOCX export")
    doc = Document()
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)
    doc.save(output)


def export_to_docx(text: str) -> bytes:
    """Return a DOCX bytes object for the given markdown text."""
    with io.BytesIO() as buf:
        _render_markdown_to_docx(text, buf)
        return buf.getvalue()


def draft_from_template(template_path: _pl.Path, facts: str, model_name: str | None = None) -> Tuple[str, Dict[str, Tuple[int, int]]]:
    """Fill a pleading template with user facts and AI-generated sections."""
    template = load_template(template_path)
    model_name = model_name or GEMINI_MODEL_DEFAULT
    model, err = get_gemini_model(model_name)
    if not model:
        return f"Error: {err}", {}

    tokens: Dict[str, Tuple[int, int]] = {}

    template = template.replace("{{FACTS}}", facts.strip())

    for match in re.findall(r"{{AI_([A-Za-z0-9_]+)}}", template):
        prompt = f"Write the {match.replace('_', ' ')} section of a UK pleading. Facts:\n{facts.strip()}"
        gen_cfg = genai_types.GenerationConfig(temperature=0.3, max_output_tokens=512) if genai_types else None
        ai_text, p_tok, c_tok = _gemini_generate_content_with_retry_and_tokens(
            model, [prompt], gen_cfg, DEFAULT_GEMINI_SAFETY_SETTINGS, "Pleading", match
        )
        template = template.replace(f"{{{{AI_{match}}}}}", ai_text)
        tokens[match] = (p_tok, c_tok)

    return template, tokens
