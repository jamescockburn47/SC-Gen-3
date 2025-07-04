# app_utils.py

import json
import re
import logging
from typing import Tuple, Optional, List, Dict, Union, Callable, Any
import io
import pathlib as _pl
from urllib.parse import quote_plus

# --- Core Dependencies ---
# Pylance will report these as unresolved if not in the environment.
# Ensure they are in your requirements.txt and installed via pip.

# Get the logger from the config first, as it's used in the try-except blocks below.
# This assumes config.py can be imported without issues.
try:
    from config import logger
except ImportError:
    # Fallback basic logger if config.py or its logger isn't available
    logging.basicConfig(level=logging.INFO)
    logger = logging.getLogger(__name__)
    logger.warning("Could not import logger from config.py, using fallback basicConfig for app_utils.")

try:
    import requests 
    # Alias for requests.exceptions.RequestException to be used in except blocks
    RequestException = requests.exceptions.RequestException
except ImportError:
    logger.warning("The 'requests' library is not installed. Network-dependent features will fail in app_utils.")
    requests = None # type: ignore
    # Fallback to generic Exception for type hinting and except blocks if requests is missing
    RequestException = Exception # type: ignore

try:
    from bs4 import BeautifulSoup
except ImportError:
    logger.warning("The 'BeautifulSoup4' (bs4) library is not installed. HTML parsing features will fail in app_utils.")
    BeautifulSoup = None # type: ignore

try:
    from PyPDF2 import PdfReader
except ImportError:
    logger.warning("The 'PyPDF2' library is not installed. PDF processing features will be limited in app_utils.")
    PdfReader = None # type: ignore

try:
    from docx import Document
except ImportError:
    logger.warning("The 'python-docx' (docx) library is not installed. DOCX processing features will fail in app_utils.")
    Document = None # type: ignore

# --- Imports from project ---
from config import (
    get_openai_client,
    get_ch_session,
    PROTO_TEXT_FALLBACK,
    LOADED_PROTO_TEXT,
    # logger is already imported/handled above
    MIN_MEANINGFUL_TEXT_LEN,
    OPENAI_MODEL_DEFAULT,
    get_gemini_model,
    GEMINI_MODEL_DEFAULT,
    GEMINI_API_KEY,
)

try:
    import google_drive_utils
except ImportError:  # pragma: no cover - optional
    google_drive_utils = None  # type: ignore

# --- Optional AWS SDK imports ---
try:
    import boto3
    from botocore.exceptions import NoCredentialsError, PartialCredentialsError, ClientError as BotoClientError
    BOTO3_AVAILABLE = True
except ImportError:
    boto3 = None # type: ignore
    # Define dummy exceptions if botocore is not available, so type checking doesn't fail later
    # This helps Pylance parse the file even if the library isn't installed.
    class NoCredentialsError(Exception): pass # type: ignore
    class PartialCredentialsError(Exception): pass # type: ignore
    class BotoClientError(Exception): pass # type: ignore
    BOTO3_AVAILABLE = False
    logger.info("boto3 library or botocore.exceptions not found. AWS-dependent features will be limited in app_utils.")

try:
    # Commenting out AWS imports to remove dependencies and log messages
    # from aws_textract_utils import perform_textract_ocr as aws_perform_textract_ocr 
    # from aws_textract_utils import TEXTRACT_AVAILABLE as AWS_TEXTRACT_MODULE_AVAILABLE
    aws_perform_textract_ocr = None # type: ignore
    AWS_TEXTRACT_MODULE_AVAILABLE = False  # AWS functionality disabled
    # logger.info("Successfully imported from aws_textract_utils for app_utils.")
except ImportError:
    # logger.warning("Could not import from aws_textract_utils in app_utils. Textract fallback for user uploads will be disabled.")
    aws_perform_textract_ocr = None # type: ignore
    AWS_TEXTRACT_MODULE_AVAILABLE = False


def _word_cap(word_count: int, level: int = 2) -> int:
    """Determine an approximate word cap based on text length and detail level."""
    if word_count <= 2000:
        base = max(150, int(word_count * 0.15))
    elif word_count <= 10000:
        base = 300
    else:
        base = min(500, int(word_count * 0.05))

    scale = {1: 0.5, 2: 1.0, 3: 1.5}.get(level, 1.0)
    cap = int(base * scale)
    return max(50, min(1000, cap))

def summarise_with_title(
    text: str,
    topic: str,
    detail_level: int = 2,
) -> Tuple[str, str]:
    """Generate a short title and summary for UI display of uploaded documents.

    Uses a cost-effective OpenAI model and ``LOADED_PROTO_TEXT`` from ``config``.
    """
    if not text or not text.strip():
        return "Empty Content", "No text was provided for summarization."

    word_count = len(text.split())
    summary_word_cap = _word_cap(word_count, detail_level)
    text_to_summarise = text[:15000]
    max_tokens_for_response = int(summary_word_cap * 2.0)
    openai_model_for_this_task = OPENAI_MODEL_DEFAULT
    openai_client = get_openai_client()

    current_protocol_text = LOADED_PROTO_TEXT
    prompt = (
        f"Return ONLY valid JSON in the format {{\"title\": \"<A concise title of less than 12 words>\", "
        f"\"summary\": \"<A summary of approximately {summary_word_cap} words, capturing the essence of the text>\"}}.\n\n"
        f"Analyze the following text:\n---\n{text_to_summarise}\n---"
    )
    raw_response_content: Optional[str] = None
    title = "Error in Summarization"
    summary = "Could not generate summary due to an issue."

    # ─── Prefer Gemini if API key available ──────────────────────────────
    if GEMINI_API_KEY:
        gemini_model_client, init_error = get_gemini_model(GEMINI_MODEL_DEFAULT)
        if gemini_model_client:
            try:
                gemini_prompt = f"{current_protocol_text}\n\n{prompt}"
                gemini_response = gemini_model_client.generate_content(gemini_prompt)
                if hasattr(gemini_response, 'text') and gemini_response.text:
                    raw_response_content = gemini_response.text
                elif hasattr(gemini_response, 'parts') and gemini_response.parts:
                    raw_response_content = ''.join(part.text for part in gemini_response.parts if hasattr(part, 'text'))
                if raw_response_content:
                    raw_response_content = raw_response_content.strip()
                    data = json.loads(raw_response_content)
                    title = str(data.get("title", "Title Missing"))
                    summary = str(data.get("summary", "Summary Missing"))
                    logger.info(
                        f"Successfully generated title/summary for topic '{topic}' using Gemini model {GEMINI_MODEL_DEFAULT}."
                    )
                    return title, summary
                else:
                    logger.error(
                        f"Empty content in Gemini response for summarise_with_title (topic: {topic}). Falling back to OpenAI."
                    )
            except json.JSONDecodeError as e_json:
                raw_preview = str(raw_response_content)[:200] if raw_response_content is not None else "N/A"
                logger.error(
                    f"JSONDecodeError in summarise_with_title using Gemini model {GEMINI_MODEL_DEFAULT}: {e_json}. Raw response: {raw_preview}..."
                )
            except Exception as e:
                raw_preview = str(raw_response_content)[:200] if raw_response_content is not None else "N/A"
                logger.error(
                    f"Exception in summarise_with_title using Gemini model {GEMINI_MODEL_DEFAULT} for topic {topic}: {e}. Raw response: {raw_preview}...",
                    exc_info=True,
                )
        else:
            logger.error(
                f"Failed to initialize Gemini model for summarise_with_title (topic: {topic}): {init_error}. Falling back to OpenAI."
            )

    # ─── OpenAI Fallback ─────────────────────────────────────────────────
    if not openai_client:
        logger.error(
            f"OpenAI client not available for summarise_with_title (topic: {topic})."
        )
        return "Summarization Error", "OpenAI client not configured."

    try:
        response = openai_client.chat.completions.create(
            model=openai_model_for_this_task, temperature=0.2, max_tokens=max_tokens_for_response,
            messages=[ {"role": "system", "content": current_protocol_text}, {"role": "user", "content": prompt}],
            response_format={"type": "json_object"} 
        )
        if response.choices and response.choices[0].message:
            raw_response_content = response.choices[0].message.content
            if raw_response_content:
                 raw_response_content = raw_response_content.strip()
                 data = json.loads(raw_response_content)
                 title = str(data.get("title", "Title Missing"))
                 summary = str(data.get("summary", "Summary Missing"))
                 logger.info(f"Successfully generated title/summary for topic '{topic}' using {openai_model_for_this_task}.")
            else:
                logger.error(f"Empty content in AI response for summarise_with_title (topic: {topic}).")
                title, summary = "Summarization Error", "AI response was empty."
        else:
            logger.error(f"Invalid AI response structure for summarise_with_title (topic: {topic}).")
            title, summary = "Summarization Error", "AI response structure invalid."
    except json.JSONDecodeError as e_json:
        raw_preview = str(raw_response_content)[:200] if raw_response_content is not None else "N/A"
        logger.error(f"JSONDecodeError in summarise_with_title (topic: {topic}, model: {openai_model_for_this_task}): {e_json}. Raw response: {raw_preview}...")
        title, summary = "Summarization Format Error", f"Failed to parse AI response as JSON. Preview: {raw_preview[:150]}..."
    except Exception as e:
        raw_preview = str(raw_response_content)[:200] if raw_response_content is not None else "N/A"
        logger.error(f"Exception in summarise_with_title (topic: {topic}, model: {openai_model_for_this_task}): {e}. Raw response: {raw_preview}...", exc_info=True)
        first_part = str(raw_response_content).split(".")[0][:75].strip() if raw_response_content else ""
        title = first_part if first_part else f"Summarization Failed ({type(e).__name__})"
        summary = str(raw_response_content) if raw_response_content else "No response content or malformed."
    return title, summary

def fetch_url_content(url: str) -> Tuple[Optional[str], Optional[str]]:
    """Fetches and extracts text content from a URL. Returns (text, error_message)."""
    if not requests: # Check if requests library failed to import
        logger.error("Requests library not available. Cannot fetch URL content.")
        return None, "Network library (requests) not available."

    response_obj: Optional[requests.Response] = None
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        response_obj = requests.get(url, headers=headers, timeout=20)
        # Explicitly check if response_obj is None after the call, though unlikely if requests.get itself doesn't raise.
        if response_obj is None:
             logger.error(f"requests.get returned None for URL '{url}' unexpectedly.")
             return None, "Fetch URL Error: Network request returned no response object."
        response_obj.raise_for_status() 
    except RequestException as e: # This will catch specific requests exceptions like HTTPError, ConnectionError, Timeout
        logger.warning(f"Failed to fetch URL '{url}': {e}")
        return None, f"Fetch URL Error: {e.__class__.__name__} for {url}"
    except Exception as e_get: # Catch any other unexpected error during the GET request itself
        logger.error(f"Unexpected error during GET request for URL '{url}': {e_get}", exc_info=True)
        return None, f"Fetch URL Error: Unexpected error {type(e_get).__name__} for {url}"

    # If response_obj is None here, it means the try block for requests.get was exited due to an exception.
    # The return statements within the except blocks should have already handled this.
    # This check is a safeguard for Pylance if it cannot follow the control flow perfectly.
    if response_obj is None:
        logger.error(f"Response object is None after GET request block for URL '{url}'. This indicates an issue in error handling.")
        return None, "Fetch URL Error: No response object available after network attempt."

    try:
        # At this point, response_obj should be a valid requests.Response object
        try: html_content = response_obj.content.decode('utf-8')
        except UnicodeDecodeError:
            try: 
                apparent_enc = response_obj.apparent_encoding
                html_content = response_obj.content.decode(apparent_enc) if apparent_enc else response_obj.content.decode('ISO-8859-1', errors='replace')
            except (UnicodeDecodeError, TypeError): 
                html_content = response_obj.content.decode('ISO-8859-1', errors='replace')
        
        if not BeautifulSoup: # Check if BeautifulSoup library failed to import
            logger.error("BeautifulSoup library not available. Cannot parse HTML content.")
            return None, "HTML parsing library (BeautifulSoup) not available."
            
        soup = BeautifulSoup(html_content, "html.parser")
        for s_tag in soup(["script", "style", "nav", "header", "footer", "aside", "form", "button", "input", "meta", "link"]):
            s_tag.decompose()
        main_content_tags = soup.find_all(['main', 'article', 'div.content', 'div.main-content', 'div.post-body', 'section.content-section'])
        content_text = " ".join(tag.get_text(" ", strip=True) for tag in main_content_tags) if main_content_tags else ""
        if not content_text.strip() or len(content_text.split()) < 50 : 
            body_tag = soup.find("body")
            content_text = body_tag.get_text(" ", strip=True) if body_tag else soup.get_text(" ", strip=True) 
        content_text = re.sub(r'\s\s+', ' ', content_text).strip()
        content_text = re.sub(r'(\n\s*){3,}', '\n\n', content_text).strip() 
        if not content_text.strip():
            logger.info(f"No significant text extracted from URL '{url}' after parsing.")
            return None, f"No text content found at {url} after parsing."
        logger.info(f"Successfully extracted text from URL '{url}' ({len(content_text)} chars).")
        return content_text, None
    except Exception as e:
        logger.error(f"Failed to process content from URL '{url}': {e}", exc_info=True)
        return None, f"Process URL Error: {str(e)} for {url}"

def find_company_number(query: str, ch_api_key: Optional[str]) -> Tuple[Optional[str], Optional[str], Optional[Dict[str, Any]]]:
    """
    Searches Companies House for a company number.
    """
    if not requests: # Check if requests library failed to import
        logger.error("Requests library not available. Cannot search Companies House.")
        return None, "Network library (requests) not available.", None

    ch_session = get_ch_session(api_key_override=ch_api_key) 
    if not ch_session.auth: 
        return None, "Companies House API Key is missing or not configured for the session.", None
    if not query or not query.strip():
        return None, "Please enter a company name or number to search.", None

    query_cleaned = query.strip().upper()
    company_no_match = re.fullmatch(r"((SC|NI|OC|SO|R[0-9])?([0-9]{6,8})|[A-Z0-9]{8})", query_cleaned.replace(" ", ""))
    if company_no_match:
        potential_no = query_cleaned.replace(" ", "")
        if potential_no.isdigit() and len(potential_no) < 8:
            formatted_no = potential_no.zfill(8)
            if re.fullmatch(r"[0-9]{8}", formatted_no):
                 logger.info(f"Using provided/formatted company number: {formatted_no}")
                 return formatted_no, None, {"company_number": formatted_no, "title": "Input as Number (Formatted)"}
        if re.fullmatch(r"[A-Z0-9]{8}|(SC|NI|OC|SO|R[0-9]?)[0-9]{6}", potential_no): 
            logger.info(f"Using provided company number: {potential_no}")
            return potential_no, None, {"company_number": potential_no, "title": "Input as Number/Code"}

    logger.info(f"Searching Companies House for name/number: '{query}' (Cleaned: '{query_cleaned}')")
    search_url = "https://api.company-information.service.gov.uk/search/companies"
    params = {'q': query_cleaned, 'items_per_page': 5} 
    response_obj: Optional[requests.Response] = None 
    try:
        response_obj = ch_session.get(search_url, params=params, timeout=15)
        if response_obj is None: # Should not happen
             logger.error(f"requests.get returned None for CH search '{query}' unexpectedly.")
             return None, "CH Search Error: Network request returned no response object.", None
        response_obj.raise_for_status()
        data = response_obj.json() # Pylance might still warn if response_obj can be None from its perspective
        items: List[Dict[str, Any]] = data.get("items", []) 
    except RequestException as e: # Use the aliased RequestException
        logger.error(f"Companies House search API error for '{query}': {e}", exc_info=True)
        return None, f"Companies House Search Error: {e}", None
    except json.JSONDecodeError as e_json:
        # Ensure response_obj is checked before accessing .text
        response_text_preview = response_obj.text[:200] if response_obj is not None else "N/A (response object was None)"
        logger.error(f"Failed to decode JSON from CH search for '{query}': {e_json}. Response text: {response_text_preview}", exc_info=True)
        return None, "Companies House Search Error: Could not parse response.", None
    except Exception as e_search: 
        logger.error(f"Unexpected error during CH search for '{query}': {e_search}", exc_info=True)
        return None, f"Companies House Search Error: Unexpected error {type(e_search).__name__}.", None
    
    # If response_obj is None here, it means the try block for requests.get was exited.
    if response_obj is None:
        logger.error(f"Response object is None after CH search block for '{query}'. This indicates an issue in error handling.")
        return None, "CH Search Error: No response object available after network attempt.", None

    if not items:
        logger.warning(f"No company found for query '{query}'.")
        return None, f"No company found for '{query}'.", None
    if query_cleaned.isalnum(): 
        for item in items:
            if item.get("company_number") == query_cleaned:
                logger.info(f"Exact company number match found via search: {item.get('title')} ({item.get('company_number')})")
                return item.get("company_number"), None, item
    first_match = items[0]
    company_number_found: Optional[str] = first_match.get("company_number") 
    company_name_found: str = first_match.get("title", "N/A") 
    if company_number_found:
        logger.info(f"Found via search (first result): {company_name_found} ({company_number_found}). Using this number.")
        return company_number_found, None, first_match
    else:
        logger.warning(f"First match for '{query}' ({company_name_found}) has no company number in API response.")
        return None, "First match found but no company number available in the result.", first_match

def extract_text_from_uploaded_file(
    file_obj: io.BytesIO,
    file_name: str,
    ocr_preference: str = "aws",
) -> Tuple[Optional[str], Optional[str]]:
    """
    Extracts text from an uploaded file object (PDF, DOCX, TXT).
    """
    file_ext = file_name.split(".")[-1].lower() if '.' in file_name else ''
    text_content: Optional[str] = None
    error_message: Optional[str] = None
    
    try:
        file_obj.seek(0) 
        if file_ext == "pdf":
            if not PdfReader: 
                logger.error(f"PyPDF2 library not available. Cannot process PDF '{file_name}' with direct extraction.")
                error_message = "PDF processing library (PyPDF2) not available."
                # No text_content from PyPDF2, so it will proceed to Textract check if error_message is not fatal
            else:
                try:
                    reader = PdfReader(file_obj) 
                    pdf_text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
                    if pdf_text_parts:
                        text_content = "\n".join(pdf_text_parts).strip()
                        logger.info(f"Successfully extracted text from PDF '{file_name}' using PyPDF2.")
                except Exception as e_pypdf: 
                    logger.warning(f"PyPDF2 failed to extract text from PDF '{file_name}': {e_pypdf}. Will try OCR if available.")
                    text_content = None # Ensure text_content is None if PyPDF2 fails

            # Condition to try Textract:
            # 1. PyPDF2 didn't produce meaningful text (text_content is None or too short)
            # AND 2. Textract is available and the function to call it is defined.
            if not text_content or len(text_content) < MIN_MEANINGFUL_TEXT_LEN:
                if ocr_preference == "google":
                    g_service = config.get_google_drive_service()
                    if g_service:
                        file_obj.seek(0)
                        pdf_bytes = file_obj.getvalue()
                        gdoc_id = google_drive_utils.upload_pdf_for_ocr(g_service, pdf_bytes, file_name)
                        if gdoc_id:
                            ocr_text = google_drive_utils.extract_text_from_google_doc(g_service, gdoc_id)
                            if ocr_text and len(ocr_text.strip()) >= MIN_MEANINGFUL_TEXT_LEN:
                                text_content = ocr_text.strip()
                                error_message = None
                                logger.info(f"Extracted text from PDF '{file_name}' using Google Drive OCR.")
                            else:
                                logger.warning(f"Google OCR for '{file_name}' returned insufficient text.")

                if (not text_content or len(text_content) < MIN_MEANINGFUL_TEXT_LEN) and AWS_TEXTRACT_MODULE_AVAILABLE and aws_perform_textract_ocr and (ocr_preference == "aws" or ocr_preference == "google"):
                    current_text_len = len(text_content) if text_content else 0
                    logger.info(f"Direct PDF extraction for '{file_name}' yielded minimal/no text (length {current_text_len}, threshold {MIN_MEANINGFUL_TEXT_LEN}). Attempting AWS Textract OCR.")
                    file_obj.seek(0)
                    pdf_bytes = file_obj.getvalue()
                    ocr_text, pages_ocrd, ocr_error = aws_perform_textract_ocr(pdf_bytes, file_name)
                    if ocr_error:
                        new_ocr_error_msg = f"Textract OCR failed for '{file_name}': {ocr_error}"
                        error_message = f"{error_message} | {new_ocr_error_msg}" if error_message else new_ocr_error_msg
                        logger.error(new_ocr_error_msg)
                    elif ocr_text:
                        text_content = ocr_text.strip()
                        error_message = None
                        logger.info(f"Successfully extracted text from PDF '{file_name}' using AWS Textract ({pages_ocrd} pages).")
                    else:
                        logger.warning(f"Textract OCR for '{file_name}' returned no text but no error. Using previous extraction if any (text_content length: {len(text_content or '')}).")
                elif text_content:
                    logger.info(f"Direct PDF extraction for '{file_name}' yielded minimal text (length {len(text_content)}), OCR not further attempted.")
                elif not error_message:
                    error_message = f"Failed to extract text from PDF '{file_name}'. No OCR method succeeded." 
                    logger.warning(error_message)
        
        elif file_ext == "docx":
            if not Document: 
                logger.error(f"python-docx library not available. Cannot process DOCX '{file_name}'.")
                error_message = "DOCX processing library (python-docx) not available."
            else:
                doc = Document(file_obj) 
                text_content = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
                logger.info(f"Successfully extracted text from DOCX: {file_name}")
        
        elif file_ext == "txt":
            try:
                text_content = file_obj.getvalue().decode("utf-8", "ignore").strip()
            except Exception as e_txt_decode_utf8:
                logger.warning(f"Could not decode TXT {file_name} as UTF-8 ({e_txt_decode_utf8}), trying fallback.")
                file_obj.seek(0) 
                try: text_content = file_obj.read().decode(errors='replace').strip()
                except Exception as e_txt_read_fallback:
                    error_message = f"Could not decode TXT file {file_name} with fallbacks: {e_txt_read_fallback}"
                    logger.error(error_message)
            if text_content and not error_message: 
                 logger.info(f"Successfully extracted text from TXT: {file_name}")
        else:
            if not error_message: # Only set this if no other error has been logged for this file
                error_message = f"Unsupported file type: '{file_ext}' for file '{file_name}'"
            logger.warning(f"Unsupported file type: '{file_ext}' for file '{file_name}' (Error: {error_message})")


        if text_content is not None and not text_content.strip(): 
            text_content = None 
            if not error_message: 
                 logger.info(f"No text content extracted from {file_name} (empty after extraction process).")
        elif text_content and not error_message: 
             logger.info(f"Extraction complete for {file_name}. Text length: {len(text_content)} chars.")
    except Exception as e: 
        logger.error(f"Error reading or processing uploaded file {file_name} (ext: {file_ext}): {e}", exc_info=True)
        error_message = f"File Read/Process Error for {file_name}: {str(e)}"
        text_content = None 
        
    return text_content, error_message


def extract_text_from_google_drive_file(
    service: Any,
    file_id: str,
    file_name: str,
    mime_type: str,
    ocr_preference: str = "google",
) -> Tuple[Optional[str], Optional[str]]:
    """Download a file from Google Drive and extract its text."""
    if not service:
        return None, "Google Drive service not available"
    if mime_type == "application/vnd.google-apps.document":
        text = google_drive_utils.extract_text_from_google_doc(service, file_id) if google_drive_utils else None
        return text, None if text else (None, "Failed to export Google Doc")
    file_bytes = google_drive_utils.download_file_bytes(service, file_id) if google_drive_utils else None
    if file_bytes is None:
        return None, f"Failed to download file {file_name}"
    return extract_text_from_uploaded_file(io.BytesIO(file_bytes), file_name, ocr_preference)


def build_consult_docx(conversation: List[str], output_path: _pl.Path) -> None:
    """Create a DOCX export of an AI consultation conversation."""
    if not Document:
        raise ImportError("python-docx library is required for DOCX export")

    doc = Document()
    doc.add_heading("AI Consultation Export", level=0)

    def _add_lines(text: str, style: str | None = None) -> None:
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith(('- ', '* ')):
                doc.add_paragraph(stripped[2:].strip(), style or "List Bullet")
            else:
                doc.add_paragraph(stripped, style=style)

    for idx, entry in enumerate(conversation, 1):
        instr = entry
        resp = ""
        m = re.search(r"Instruction:\s*(.*?)\n\n?Response[^:]*?:\s*(.*)", entry, re.S)
        if m:
            instr, resp = m.group(1).strip(), m.group(2).strip()

        doc.add_heading(f"Instruction {idx}", level=1)
        _add_lines(instr)
        if resp:
            doc.add_heading("Response", level=2)
            _add_lines(resp)

    doc.save(output_path)


def extract_legal_citations(text: str) -> List[str]:
    """Return a list of case and statute citations found in ``text``."""
    if not text:
        return []

    # More precise pattern for case citations - avoid capturing leading words
    case_pat = re.compile(r"(?:^|[^A-Za-z])([A-Z][A-Za-z0-9&.,'\-]{1,30}(?: [A-Z][A-Za-z0-9&.,'\-]{1,30}){0,4} v\.? [A-Z][A-Za-z0-9&.,'\-]{1,30}(?: [A-Za-z0-9&.,'\-]{1,30}){0,4} \[[0-9]{4}\][^A-Za-z]*)")
    # Pattern for legislation citations
    statute_pat = re.compile(r"\b([A-Z][A-Za-z0-9()\-\s]{1,50} (?:Act|Regulations) [12][0-9]{3})\b")

    citations: List[str] = []
    
    # Extract case citations
    for m in case_pat.finditer(text):
        cit = m.group(1).strip()
        # Clean up the citation
        cit = re.sub(r'^[^A-Z]*', '', cit)  # Remove any leading non-capital characters
        cit = re.sub(r'[^A-Za-z0-9\[\]]*$', '', cit)  # Remove trailing non-alphanumeric except brackets
        if cit and len(cit) > 10 and cit not in citations:  # Ensure minimum length
            citations.append(cit)
    
    # Extract statute citations
    for m in statute_pat.finditer(text):
        cit = m.group(1).strip()
        if cit not in citations:
            citations.append(cit)
    
    return citations


def verify_citations(
    citations: List[str],
    uploaded_files: List[Any],
    cache_file: Optional[_pl.Path] = None,
) -> Dict[str, bool]:
    """Check citations against uploaded files or public sources."""

    verified_cache: Dict[str, Any] = {}
    if cache_file and cache_file.exists():
        try:
            verified_cache = json.loads(cache_file.read_text())
        except Exception as e:
            logger.warning(f"Failed to read citation cache {cache_file}: {e}")

    # Re-check any cached entries that include URLs but were not yet verified
    cache_updated = False
    for cit, entry in list(verified_cache.items()):
        if isinstance(entry, dict) and entry.get("url") and not entry.get("verified"):
            text, _err = fetch_url_content(entry["url"])
            if text and cit.lower() in text.lower():
                entry["verified"] = True
                verified_cache[cit] = entry
                cache_updated = True

    uploaded_texts: List[str] = []
    for f in uploaded_files:
        try:
            txt, _ = extract_text_from_uploaded_file(io.BytesIO(f.getvalue()), f.name)
            if txt:
                uploaded_texts.append(txt.lower())
        except Exception as e:
            logger.error(f"Error extracting text from {getattr(f, 'name', '?')}: {e}")

    results: Dict[str, bool] = {}
    for cit in citations:
        cache_entry = verified_cache.get(cit)
        cached_verified = False
        if isinstance(cache_entry, bool):
            cached_verified = cache_entry
        elif isinstance(cache_entry, dict):
            cached_verified = bool(cache_entry.get("verified"))

        if cached_verified:
            results[cit] = True
            continue

        found = any(cit.lower() in t for t in uploaded_texts)

        if not found:
            # Try multiple search strategies
            search_bases = [
                "https://www.casemine.com/search?q=",
                "https://www.bailii.org/cgi-bin/form.pl?db=&ctl=Search&method=boolean&query=",
            ]
            
            for base in search_bases:
                url = base + quote_plus(cit)
                text, err = fetch_url_content(url)
                if text:
                    # More flexible matching - check for key parts of the citation
                    citation_parts = cit.replace('[', '').replace(']', '').split()
                    matches = sum(1 for part in citation_parts if len(part) > 2 and part.lower() in text.lower())
                    # Consider verified if most parts of the citation are found
                    if matches >= len(citation_parts) * 0.6:  # 60% of parts found
                        found = True
                        break

        results[cit] = found
        if found:
            cache_updated = True
            if isinstance(cache_entry, dict):
                cache_entry["verified"] = True
                verified_cache[cit] = cache_entry
            else:
                verified_cache[cit] = True

    if cache_file:
        try:
            if cache_updated:
                cache_file.write_text(json.dumps(verified_cache, indent=2))
        except Exception as e:
            logger.warning(f"Failed to update citation cache {cache_file}: {e}")

    return results
